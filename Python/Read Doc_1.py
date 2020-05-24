import docx
def ReadingTestDocuments(filename):
    doc=docx.Document(filename)

    completeText = []
    for paragraph in doc.paragraphs:
        completeText.append(paragraph.text)

    return "\n" .join(completeText)

print(ReadingTestDocuments("File.docx"))
'''
from docx import Document
wordDoc = Document('File.docx')
for table in wordDoc.tables:
    for row in table.rows:
        for cell in row.cells:
            print (cell.text)
            

'''
'''
from docx import Document
 
document = Document('filename.docx')
tbl = document.tables[1]
for rw in tbl.rows:
    if rw.cells[0].text.startswith('Sr.No'):
        print(rw.cells[3].text)
'''
from docx import Document
document = Document('File.docx')
tables = document.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)
